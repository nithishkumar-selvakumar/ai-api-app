from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import (
    CSVLoader,
    JSONLoader,
    PyPDFLoader,
    UnstructuredExcelLoader,
)
from langchain_core.documents import Document


def load_document(file_path: Path) -> list[Document]:
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return PyPDFLoader(str(file_path)).load()

    if extension == ".csv":
        return CSVLoader(str(file_path)).load()

    if extension in {".xlsx", ".xls"}:
        return UnstructuredExcelLoader(
            str(file_path)
        ).load()

    if extension == ".json":
        return JSONLoader(
            file_path=str(file_path),
            jq_schema=".",
            text_content=False,
        ).load()

    if extension == ".docx":
        return _load_docx(file_path)

    raise ValueError(
        f"Unsupported document type: {extension}"
    )


def _load_docx(file_path: Path) -> list[Document]:
    docx = DocxDocument(str(file_path))

    text = "\n".join(
        paragraph.text
        for paragraph in docx.paragraphs
        if paragraph.text.strip()
    )

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": str(file_path),
                "file_type": "docx",
            },
        )
    ]